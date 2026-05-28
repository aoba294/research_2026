from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\大谷　青羽\Desktop\進捗報告_2026\進捗報告_2026-05-25_転倒検知モデル.docx")


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_paragraph(paragraph, size=11, bold=False, color=None):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        size = 16
        color = BLUE
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        size = 13
        color = BLUE if level == 2 else DARK_BLUE
    style_paragraph(p, size=size, bold=True, color=color)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    style_paragraph(p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    style_paragraph(p)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, headers):
        cell.text = text
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
            for p in cell.paragraphs:
                style_paragraph(p)
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("進捗報告: 仮想環境における転倒検知モデルの作成")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    meta = doc.add_paragraph("日付: 2026年5月25日")
    meta.paragraph_format.space_after = Pt(12)
    style_paragraph(meta, size=10, color=RGBColor(0x55, 0x55, 0x55))

    add_heading(doc, "1. 本日の目的", 1)
    add_body(
        doc,
        "これまでの生成データは、転倒時にz方向の特徴量だけで判定できるほど単純であり、"
        "実環境での転倒検知可能性を評価するには不十分であった。そこで本日は、仮想環境の段階で"
        "転倒検知が成立するモデルを作ることを目標に、生成データの多様化、評価指標の見直し、"
        "5フレーム時系列モデルの作成を行った。",
    )

    add_heading(doc, "2. 実施内容", 1)
    for item in [
        "転倒時のLiDAR点群を単一パターンから複数パターンに変更した。",
        "通常行動にも前かがみ、座る途中、一部遮蔽などの転倒に似た動作を追加した。",
        "全体accuracyだけでなく、fall recall、fall precision、false alarm rate、confusion matrixを出力する評価に変更した。",
        "転倒データを学習用に多めに生成し、評価用データは現実寄りの低頻度転倒のまま残した。",
        "1フレーム判定ではなく、前後2フレームを含む5フレーム窓で転倒を判定するモデルを作成した。",
        "未知施設評価として、部屋サイズ、ノイズ、遮蔽、ベッド位置が異なる施設データを用意した。",
        "モデルを保存し、仮想環境内で再利用できる形にした。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 主な結果", 1)
    add_body(
        doc,
        "最終的には、5フレームのLiDAR変化とベッド圧変化を用いたweighted logistic regressionにより、"
        "未知施設評価でも転倒イベント単位で検知可能な水準まで改善した。",
    )

    add_table(
        doc,
        ["評価単位", "Accuracy", "Fall recall", "Fall precision", "False alarm"],
        [
            ["フレーム単位", "0.974", "0.782", "0.470", "0.021"],
            ["イベント単位", "-", "0.921", "0.442", "-"],
        ],
        [1800, 1500, 1800, 2000, 2260],
        header_fill=LIGHT_BLUE,
    )

    add_body(
        doc,
        "イベント単位では、未知施設において89件の転倒イベント中82件を検出した。"
        "一方で、アラート199件中111件は誤報であり、今後はprecision改善が課題である。",
    )

    add_heading(doc, "4. 重要な気づき", 1)
    add_table(
        doc,
        ["観点", "内容"],
        [
            [
                "accuracyの限界",
                "転倒が少ないため、全体accuracyだけでは転倒検知性能を正しく評価できない。"
                "実際に、単純なしきい値ではaccuracyが高くても転倒を見逃すケースがあった。",
            ],
            [
                "5フレーム化の効果",
                "転倒は姿勢そのものよりも、急な高さ低下や床付近点の増加などの時間変化として捉える方が有効であった。",
            ],
            [
                "転倒エピソードの重要性",
                "生成データで転倒が1フレームだけで終わると、実際の転倒検知タスクとして不自然になる。"
                "そのため、転倒後の床上状態が数フレーム継続するように生成過程を修正した。",
            ],
            [
                "未知施設評価",
                "同じ生成条件だけで評価すると性能が高く出やすい。施設差を入れた未知施設評価により、"
                "より現実に近い課題が見えるようになった。",
            ],
        ],
        [2000, 7360],
    )

    add_heading(doc, "5. 作成・変更したファイル", 1)
    add_table(
        doc,
        ["ファイル", "内容"],
        [
            [
                "heterosense/_core/_observation_model.py",
                "転倒点群の多様化、通常行動のhard negative追加、5フレーム特徴用の観測多様性強化。",
            ],
            [
                "heterosense/_core/_behavior_model.py",
                "転倒イベントが数フレーム継続するようにABNORMAL episode durationを導入。",
            ],
            [
                "experiments/check_data.py",
                "評価指標、5フレーム特徴、未知施設評価、簡易ML比較を追加。",
            ],
            [
                "experiments/train_virtual_fall_detector.py",
                "5フレーム転倒検知モデルの訓練、しきい値探索、イベント単位評価、モデル保存を実装。",
            ],
            [
                "results/virtual_fall_detector_5frame.npz",
                "仮想環境用の転倒検知モデルとして保存したファイル。",
            ],
        ],
        [3300, 6060],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "6. 今後の課題", 1)
    for item in [
        "誤報を減らすため、しゃがむ、床の物を拾う、ベッドに横になる、椅子に勢いよく座るなどのhard negativeをさらに増やす。",
        "アラート単位のprecisionを改善するため、連続検知条件、cooldown、しきい値を追加調整する。",
        "施設ごとの通常時平均との差分を特徴量に入れ、未知施設への汎化性能を高める。",
        "将来的にはRandomForestや時系列モデルなど、より表現力のある分類器との比較も行う。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. 現時点の結論", 1)
    add_body(
        doc,
        "本日の作業により、仮想環境内では転倒イベントの約92%を検知できるモデルを作成できた。"
        "ただし、現実環境では性能が下がる可能性が高いため、今後は通常行動の多様化と誤報削減を中心に改善する必要がある。",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("進捗報告 2026-05-25")
    style_paragraph(footer, size=9, color=RGBColor(0x66, 0x66, 0x66))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
