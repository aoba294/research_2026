from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\大谷　青羽\Desktop\進捗報告_2026\進捗報告_2026-05-28_1D-CNN施設補正.docx")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_paragraph(paragraph, size=11, bold=False, color=None):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    style_paragraph(p, size=16 if level == 1 else 13, bold=True, color=BLUE if level <= 2 else DARK_BLUE)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    style_paragraph(p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    style_paragraph(p)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
            for p in cell.paragraphs:
                style_paragraph(p)
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("進捗報告: 1D-CNNと施設補正による転倒検知性能評価")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    meta = doc.add_paragraph("日付: 2026年5月28日")
    style_paragraph(meta, size=10, color=GRAY)

    add_heading(doc, "1. 本日の目的")
    add_body(
        doc,
        "最終目標は実際の介護施設で転倒を検知できるモデルを作ることである。"
        "本日は、これまでの5フレーム特徴量モデルを発展させ、NNモデルで施設補正の効果を確認し、"
        "さらに1D-CNNによる時系列処理と点群の動きの多様化を行った。",
    )

    add_heading(doc, "2. 実施内容")
    for item in [
        "5フレーム特徴量を入力するMLP型NNを実装し、施設補正なし/ありで比較した。",
        "5フレームの時系列を直接入力する1D-CNNを実装した。",
        "施設ごとの通常時平均との差分を特徴量に加え、未知施設での誤報削減効果を確認した。",
        "転倒イベント内の点群の動きを多様化した。具体的には、転倒方向、転倒後の静止、寝返り、手足運動、起き上がり試行、床上移動を追加した。",
        "高すぎる結果を疑い、ノイズ・遮蔽・部屋形状・ベッド位置を厳しくしたストレス評価セットを追加した。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 主な結果")
    add_table(
        doc,
        ["モデル", "評価条件", "Event recall", "Alert precision", "False alerts"],
        [
            ["MLP NN 補正なし", "未知施設", "0.978", "0.354", "173"],
            ["MLP NN 補正あり", "未知施設", "0.910", "0.701", "38"],
            ["1D-CNN 補正なし", "未知施設", "0.904", "0.752", "30"],
            ["1D-CNN 補正あり", "未知施設", "0.936", "0.847", "17"],
            ["1D-CNN 補正あり", "厳しい評価", "0.820", "0.918", "7"],
        ],
        [2100, 1600, 1700, 1900, 2060],
        header_fill=LIGHT_BLUE,
    )

    add_body(
        doc,
        "1D-CNNに施設補正を入れることで、未知施設評価においてevent recall 0.936、"
        "alert precision 0.847となった。MLPよりも時系列変化を扱いやすく、誤報を抑えつつ転倒検知率を維持できた。",
    )

    add_heading(doc, "4. 考察")
    add_table(
        doc,
        ["観点", "内容"],
        [
            [
                "施設補正の効果",
                "施設ごとの通常時平均との差分を入れることで、施設固有のセンサ位置、部屋サイズ、ノイズの影響を軽減できた。",
            ],
            [
                "1D-CNNの効果",
                "5フレームの時間変化を直接扱うことで、単発の姿勢ではなく、転倒前後の動きとして判定できるようになった。",
            ],
            [
                "点群生成の課題",
                "高すぎる精度は生成器のクセを学習している可能性があるため、点群の時間的な動きを多様化した。",
            ],
            [
                "厳しい評価の必要性",
                "ストレス評価では検知率が下がったため、仮想環境内の高精度をそのまま現実環境に適用できるとは言えない。",
            ],
        ],
        [1900, 7460],
    )

    add_heading(doc, "5. 作成・変更したファイル")
    add_table(
        doc,
        ["ファイル", "内容"],
        [
            ["experiments/train_nn_fall_detector.py", "MLP型NNで施設補正なし/ありを比較するスクリプト。"],
            ["experiments/train_cnn_fall_detector.py", "1D-CNNで5フレーム時系列を学習し、施設補正の効果を比較するスクリプト。"],
            ["experiments/stress_test_cnn_detector.py", "保存済み1D-CNNを厳しい施設条件で評価するスクリプト。"],
            ["heterosense/_core/_behavior_model.py", "転倒イベントごとの動きパターンと転倒方向を追加。"],
            ["heterosense/_core/_observation_model.py", "転倒後の点群動作を多様化。"],
            ["results/virtual_fall_detector_cnn_5frame_baseline.npz", "施設補正あり1D-CNNモデル。"],
        ],
        [3300, 6060],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "6. 次にやること")
    for item in [
        "複数seedで1D-CNNの平均性能と標準偏差を確認する。",
        "厳しい評価セットを学習にも混ぜ、ストレス条件でのevent recallを改善する。",
        "実施設の通常行動データを少量でも取得し、施設ごとの通常平均との差分の実用性を確認する。",
        "誤報が残るケースを分類し、しゃがむ、介助動作、ベッド動作などのhard negativeを追加する。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. 現時点の結論")
    add_body(
        doc,
        "仮想環境では、1D-CNNと施設ごとの通常時平均との差分特徴を組み合わせることで、"
        "未知施設に対しても転倒検知性能が改善した。ただし、厳しい施設条件では性能が低下するため、"
        "今後は生成点群の多様化と実データによる検証が必要である。",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("進捗報告 2026-05-28")
    style_paragraph(footer, size=9, color=GRAY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
